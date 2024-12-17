import os
import math
import argparse
import torch
import torch.optim as optim
import torch.optim.lr_scheduler as lr_scheduler
from torch.utils.tensorboard import SummaryWriter
from torchvision import transforms, utils
from my_dataset import MyDataSet
from torch.utils.data import Dataset, DataLoader
from new_model import vit_base_patch16_224_in21k as create_model  # 导入预训练模型
# from RMT import RMT_T as create_model
from utils1 import read_split_data, train_one_epoch, evaluate

import pandas as pd
import xlwt
import xlrd
from skimage import transform
import numpy as np
import matplotlib.pyplot as plt  # 新增
os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"
plt.ion()

os.environ['CUDA_VISIBLE_DEVICES']='0'

import warnings
warnings.filterwarnings("ignore")

# torch.backends.cudnn.benchmark=True

book = xlwt.Workbook(encoding='utf-8')  # 创建Workbook，相当于创建Excel
# 创建sheet，Sheet1为表的名字，cell_overwrite_ok为是否覆盖单元格
sheet2 = book.add_sheet(u'Train_data', cell_overwrite_ok=True)
# 向表中添加数据
sheet2.write(0, 0, 'epoch')
sheet2.write(0, 1, 'Train_Loss')
sheet2.write(0, 2, 'Train_Acc')
sheet2.write(0, 3, 'Val_Loss')
sheet2.write(0, 4, 'Val_Acc')
sheet2.write(0, 5, 'lr')
sheet2.write(0, 6, 'Best val Acc')


import csv

from docx import Document
import pandas as pd


def save_model_layers_and_params_to_word(model, epoch, output_file):
    document = Document()
    document.add_heading(f'Epoch {epoch} - Model Layers and Parameters', level=1)

    for name, param in model.named_parameters():
        document.add_heading(name, level=2)
        document.add_paragraph(f"Requires Grad: {param.requires_grad}")
        document.add_paragraph(f"Shape: {list(param.shape)}")

    document.save(output_file)


def save_model_layers_and_params_to_excel(model, epoch, output_file):
    layers_info = []

    for name, param in model.named_parameters():
        layers_info.append({
            "Layer Name": name,
            "Requires Grad": param.requires_grad,
            "Shape": str(list(param.shape))
        })

    df = pd.DataFrame(layers_info)
    df.to_excel(output_file, index=False)


def read_csv(file_path):
    data = []
    with open(file_path, 'r') as file:
        csv_reader = csv.reader(file)
        for row in csv_reader:
            data.append(row)
    return data


# 三个数据变换类的定义
class Rescale(object):
    """Rescale the image in a sample to a given size.
        对图片和控制点的坐标进行尺度变换。
    Args:
        output_size (tuple or int): Desired output size. If tuple, output is
            matched to output_size. If int, smaller of image edges is matched
            to output_size keeping aspect ratio the same.
    """
    def __init__(self, output_size):
        assert isinstance(output_size,  (int, tuple))
        self.output_size = output_size

    def __call__(self, sample):
        image, face_point, point, lh_point, rh_point, label = sample['frame'], sample['face_point'], sample['point'], \
                                                              sample['left_hand_point'], sample['right_hand_point'], \
                                                              sample['Y']

        h, w = image.shape[:2]
        if isinstance(self.output_size, int):
            if h > w:
                new_h, new_w = self.output_size * h / w, self.output_size
            else:
                new_h, new_w = self.output_size, self.output_size * w / h
        else:
            new_h, new_w = self.output_size

        new_h, new_w = int(new_h), int(new_w)
        img = transform.resize(image, (new_h, new_w))

        return {'frame': img, 'face_point': face_point, 'point': point,
                'left_hand_point': lh_point, 'right_hand_point': rh_point, 'Y': label}


class RandomCrop(object):
    """Crop randomly the image in a sample.
        图片进行规定尺寸的随机裁剪。控制点的坐标相应平移
    Args:
        output_size (tuple or int): Desired output size. If int, square crop
            is made.
    """
    def __init__(self, output_size):
        assert isinstance(output_size, (int, tuple))
        if isinstance(output_size, int):
            self.output_size = (output_size, output_size)
        else:
            assert len(output_size) == 2
            self.output_size = output_size

    def __call__(self, sample):
        image, face_point, point, lh_point, rh_point, label = sample['frame'], sample['face_point'], sample['point'], \
                                                              sample['left_hand_point'], sample['right_hand_point'], \
                                                              sample['Y']
        # 对图片
        h, w = image.shape[:2]
        new_h, new_w = self.output_size
        top = np.random.randint(0, h - new_h)
        left = np.random.randint(0, w - new_w)
        image = image[top: top + new_h,
                left: left + new_w]
        # # 对坐标
        # scaler = {}
        # feats = [face_point, point, lh_point, rh_point]
        # for x in feats:
        #     # 获得对应属性下的值并堆叠
        #     all_data = np.vstack(x)
        #     scaler[x] = MinMaxScaler()
        #     scaler[x].fit(all_data)
        #
        # rh_point = [scaler['hands_right'].transform(x) for x in rh_point]
        # lh_point = [scaler['hands_left'].transform(x) for x in lh_point]
        # point = [scaler['bodies'].transform(x) for x in point]
        # face_point = [scaler['face'].transform(x) for x in face_point]

        return {'frame': image, 'face_point': face_point, 'point': point,
                'left_hand_point': lh_point, 'right_hand_point': rh_point, 'Y': label}


class Center(object):
    """Crop randomly the image in a sample.
        图片进行规定尺寸的随机裁剪。控制点的坐标相应平移
    Args:
        output_size (tuple or int): Desired output size. If int, square crop
            is made.
    """
    def __init__(self, output_size):
        assert isinstance(output_size, (int, tuple))
        if isinstance(output_size, int):
            self.output_size = (output_size, output_size)
        else:
            assert len(output_size) == 2
            self.output_size = output_size

    def __call__(self, sample):
        image, face_point, point, lh_point, rh_point, label = sample['frame'], sample['face_point'], sample['point'], \
                                                              sample['left_hand_point'], sample['right_hand_point'], \
                                                              sample['Y']
        # 对图片
        h, w = image.shape[:2]
        new_h, new_w = self.output_size
        top = int((h - new_h + 1) * 0.5)
        left = int((w - new_w + 1) * 0.5)
        image = image[top: top + new_h,
                left: left + new_w]

        return {'frame': image, 'face_point': face_point, 'point': point,
                'left_hand_point': lh_point, 'right_hand_point': rh_point, 'Y': label}


class ToTensor(object):
    """Convert ndarrays in sample to Tensors.
        numpy数组到tensor的变化，另外还有维度的变化。
    """
    def __call__(self, sample):
        image, face_point, point, lh_point, rh_point, label = sample['frame'], sample['face_point'], sample['point'], \
                                                              sample['left_hand_point'], sample['right_hand_point'], \
                                                              sample['Y']
        # swap color axis because
        # numpy image: H x W x C
        # torch image: C X H X W
        image = image.transpose((2, 0, 1))
        return {'frame': torch.from_numpy(image),
                'face_point': torch.from_numpy(face_point),
                'point': torch.from_numpy(point),
                'left_hand_point': torch.from_numpy(lh_point),
                'right_hand_point': torch.from_numpy(rh_point),
                'Y': torch.from_numpy(np.array(label))
                }


def main(args):
    best_acc = 0

    # device = torch.device(args.device if torch.cuda.is_available() else "cpu")
    device = 'cuda'
    if os.path.exists("./weights") is False:
        os.makedirs("./weights")

    tb_writer = SummaryWriter()
    train_videos_path, train_labels, train_points_path, val_videos_path, val_labels, val_points_path = \
        read_split_data(args.data_path)

    # 数据预处理
    # data_transform = {
    #     "train": transforms.Compose([transforms.RandomResizedCrop(224),
    #                                  transforms.RandomHorizontalFlip(),
    #                                  transforms.ToTensor(),
    #                                  transforms.Normalize([0.5, 0.5, 0.5], [0.5, 0.5, 0.5])]),
    #     "val": transforms.Compose([transforms.Resize(256),
    #                                transforms.CenterCrop(224),
    #                                transforms.ToTensor(),
    #                                transforms.Normalize([0.5, 0.5, 0.5], [0.5, 0.5, 0.5])])}

    transformed_train_dataset = MyDataSet(videos_path=train_videos_path,
                                          videos_class=train_labels,
                                          points_path=train_points_path,
                                          transform=transforms.Compose([
                                                       Rescale(256),
                                                       RandomCrop(224),
                                                       ToTensor()
                                                   ]))

    transformed_val_dataset = MyDataSet(videos_path=val_videos_path,
                                        videos_class=val_labels,
                                        points_path=val_points_path,
                                        transform=transforms.Compose([
                                                       Rescale(256),
                                                       Center(224),
                                                       ToTensor()
                                                   ]))

# -==================================================
    # 实例化训练数据集
    # train_dataset = MyDataSet(images_path=train_images_path,
    #                           images_class=train_label,
    #                           points_path=train_points_path,
    #                           transform=data_transform["train"])
    #
    # # 实例化验证数据集
    # val_dataset = MyDataSet(images_path=val_images_path,
    #                         images_class=val_label,
    #                         points_path=val_points_path,
    #                         transform=data_transform["val"])

    batch_size = args.batch_size
    nw = min([os.cpu_count(), batch_size if batch_size > 1 else 0, 8])  # number of workers
    print("nw:", nw)
    # nw = 0
    print('Using {} dataloader workers every process'.format(nw))

    # ==============================================
    train_dataloader = DataLoader(transformed_train_dataset,
                                  batch_size=batch_size,
                                  shuffle=True,
                                  num_workers=nw,
                                  collate_fn=transformed_train_dataset.collate_fn)

    val_dataloader = DataLoader(transformed_val_dataset,
                                batch_size=batch_size,
                                shuffle=False,
                                num_workers=nw,
                                collate_fn=transformed_val_dataset.collate_fn)

    # ======================================== Model Vit =========================================
    # 初始化模型（分类类别个数，logits设为否，这里如果设为是，就会只训练这里，其他的被冻结）
    # 分别为不同模型导入参数，不冻结权重
    model_vit = create_model(num_classes=11).to(device)

    # TODO
    # if args.weights == "":
    #     assert os.path.exists(args.weights), "weights file: '{}' not exist.".format(args.weights)
    # else:
    #     weights_dict = torch.load(args.weights, map_location=device)
    #     """
    #     TODO 复制空间block参数给时间
    #     """
    #     # 删除不需要的权重
    #     del_keys = ['head.weight', 'head.bias', 'head_time.weight', 'head_time.bias']
    #     # del_model_keys = model_vit.load_state_dict(weights_dict, strict=False)[1]
    #     for k in del_keys:
    #         del weights_dict[k]
    #     # for m in del_model_keys:
    #     #     del weights_dict[m]
    #     print(model_vit.load_state_dict(weights_dict, strict=False))

    # if args.freeze_layers:
    #     miss_list = ['cls_token', 'pos_embed',  'patch_embed.proj.weight', 'patch_embed.proj.bias', 'blocks.0.norm1.weight', 'blocks.0.norm1.bias', 'blocks.0.attn.qkv.weight', 'blocks.0.attn.qkv.bias', 'blocks.0.attn.proj.weight', 'blocks.0.attn.proj.bias', 'blocks.0.norm2.weight', 'blocks.0.norm2.bias', 'blocks.0.mlp.fc1.weight', 'blocks.0.mlp.fc1.bias', 'blocks.0.mlp.fc2.weight', 'blocks.0.mlp.fc2.bias', 'blocks.1.norm1.weight', 'blocks.1.norm1.bias', 'blocks.1.attn.qkv.weight', 'blocks.1.attn.qkv.bias', 'blocks.1.attn.proj.weight', 'blocks.1.attn.proj.bias', 'blocks.1.norm2.weight', 'blocks.1.norm2.bias', 'blocks.1.mlp.fc1.weight', 'blocks.1.mlp.fc1.bias', 'blocks.1.mlp.fc2.weight', 'blocks.1.mlp.fc2.bias', 'blocks.2.norm1.weight', 'blocks.2.norm1.bias', 'blocks.2.attn.qkv.weight', 'blocks.2.attn.qkv.bias', 'blocks.2.attn.proj.weight', 'blocks.2.attn.proj.bias', 'blocks.2.norm2.weight', 'blocks.2.norm2.bias', 'blocks.2.mlp.fc1.weight', 'blocks.2.mlp.fc1.bias', 'blocks.2.mlp.fc2.weight', 'blocks.2.mlp.fc2.bias', 'blocks.3.norm1.weight', 'blocks.3.norm1.bias', 'blocks.3.attn.qkv.weight', 'blocks.3.attn.qkv.bias', 'blocks.3.attn.proj.weight', 'blocks.3.attn.proj.bias', 'blocks.3.norm2.weight', 'blocks.3.norm2.bias', 'blocks.3.mlp.fc1.weight', 'blocks.3.mlp.fc1.bias', 'blocks.3.mlp.fc2.weight', 'blocks.3.mlp.fc2.bias']
        # miss_list = model_vit.load_state_dict(weights_dict, strict=False)[0]
     #   for name, para in model_vit.named_parameters():
     #       # 除head, pre_logits外，其他权重全部冻结False表示冻结
     #       # if "head" not in name and "pre_logits" not in name:
     #       if name not in miss_list:
     #           para.requires_grad_(False)
     #           print(" freeze {}".format(name))
     #       else:
     #           para.requires_grad_(True)
     #           print("train {}".format(name))

    # if args.freeze_layers:
    #     miss_list = ['cls_token', 'pos_embed',  'patch_embed.proj.weight', 'patch_embed.proj.bias', 'blocks.0.norm1.weight', 'blocks.0.norm1.bias', 'blocks.0.attn.qkv.weight', 'blocks.0.attn.qkv.bias', 'blocks.0.attn.proj.weight', 'blocks.0.attn.proj.bias', 'blocks.0.norm2.weight', 'blocks.0.norm2.bias', 'blocks.0.mlp.fc1.weight', 'blocks.0.mlp.fc1.bias', 'blocks.0.mlp.fc2.weight', 'blocks.0.mlp.fc2.bias', 'blocks.1.norm1.weight', 'blocks.1.norm1.bias', 'blocks.1.attn.qkv.weight', 'blocks.1.attn.qkv.bias', 'blocks.1.attn.proj.weight', 'blocks.1.attn.proj.bias', 'blocks.1.norm2.weight', 'blocks.1.norm2.bias', 'blocks.1.mlp.fc1.weight', 'blocks.1.mlp.fc1.bias', 'blocks.1.mlp.fc2.weight', 'blocks.1.mlp.fc2.bias', 'blocks.2.norm1.weight', 'blocks.2.norm1.bias', 'blocks.2.attn.qkv.weight', 'blocks.2.attn.qkv.bias', 'blocks.2.attn.proj.weight', 'blocks.2.attn.proj.bias', 'blocks.2.norm2.weight', 'blocks.2.norm2.bias', 'blocks.2.mlp.fc1.weight', 'blocks.2.mlp.fc1.bias', 'blocks.2.mlp.fc2.weight', 'blocks.2.mlp.fc2.bias', 'blocks.3.norm1.weight', 'blocks.3.norm1.bias', 'blocks.3.attn.qkv.weight', 'blocks.3.attn.qkv.bias', 'blocks.3.attn.proj.weight', 'blocks.3.attn.proj.bias', 'blocks.3.norm2.weight', 'blocks.3.norm2.bias', 'blocks.3.mlp.fc1.weight', 'blocks.3.mlp.fc1.bias', 'blocks.3.mlp.fc2.weight', 'blocks.3.mlp.fc2.bias']
    #     # miss_list = model_vit.load_state_dict(weights_dict, strict=False)[0]
    #     for name, para in model_vit.named_parameters():
    #         # 除head, pre_logits外，其他权重全部冻结False表示冻结
    #         # if "head" not in name and "pre_logits" not in name:
    #         # if name not in miss_list:
    #         if name in miss_list:   # 1.27测试
    #             para.requires_grad_(False)
    #             print("freeze {}".format(name))
    #
    #         else:
    #             para.requires_grad_(True)
    #             print("training {}".format(name))
    if args.freeze_layers:
        miss_list = [
            'Resnet.module.conv1.weight', 'Resnet.module.bn1.weight', 'Resnet.module.bn1.bias',
            'Resnet.module.layer1.0.conv1.weight', 'Resnet.module.layer1.0.bn1.weight',
            'Resnet.module.layer1.0.bn1.bias',
            'Resnet.module.layer1.0.conv2.weight', 'Resnet.module.layer1.0.bn2.weight',
            'Resnet.module.layer1.0.bn2.bias',
            'Resnet.module.layer1.0.conv3.weight', 'Resnet.module.layer1.0.bn3.weight',
            'Resnet.module.layer1.0.bn3.bias',
            'Resnet.module.layer1.0.downsample.0.weight', 'Resnet.module.layer1.0.downsample.1.weight',
            'Resnet.module.layer1.0.downsample.1.bias',
            'Resnet.module.layer1.1.conv1.weight', 'Resnet.module.layer1.1.bn1.weight',
            'Resnet.module.layer1.1.bn1.bias',
            'Resnet.module.layer1.1.conv2.weight', 'Resnet.module.layer1.1.bn2.weight',
            'Resnet.module.layer1.1.bn2.bias',
            'Resnet.module.layer1.1.conv3.weight', 'Resnet.module.layer1.1.bn3.weight',
            'Resnet.module.layer1.1.bn3.bias',
            'Resnet.module.layer1.2.conv1.weight', 'Resnet.module.layer1.2.bn1.weight',
            'Resnet.module.layer1.2.bn1.bias',
            'Resnet.module.layer1.2.conv2.weight', 'Resnet.module.layer1.2.bn2.weight',
            'Resnet.module.layer1.2.bn2.bias',
            'Resnet.module.layer1.2.conv3.weight', 'Resnet.module.layer1.2.bn3.weight',
            'Resnet.module.layer1.2.bn3.bias',
            'Resnet.module.layer2.0.conv1.weight', 'Resnet.module.layer2.0.bn1.weight',
            'Resnet.module.layer2.0.bn1.bias',
            'Resnet.module.layer2.0.conv2.weight', 'Resnet.module.layer2.0.bn2.weight',
            'Resnet.module.layer2.0.bn2.bias',
            'Resnet.module.layer2.0.conv3.weight', 'Resnet.module.layer2.0.bn3.weight',
            'Resnet.module.layer2.0.bn3.bias',
            'Resnet.module.layer2.0.downsample.0.weight', 'Resnet.module.layer2.0.downsample.1.weight',
            'Resnet.module.layer2.0.downsample.1.bias',
            'Resnet.module.layer2.1.conv1.weight', 'Resnet.module.layer2.1.bn1.weight',
            'Resnet.module.layer2.1.bn1.bias',
            'Resnet.module.layer2.1.conv2.weight', 'Resnet.module.layer2.1.bn2.weight',
            'Resnet.module.layer2.1.bn2.bias',
            'Resnet.module.layer2.1.conv3.weight', 'Resnet.module.layer2.1.bn3.weight',
            'Resnet.module.layer2.1.bn3.bias',
            'Resnet.module.layer2.2.conv1.weight', 'Resnet.module.layer2.2.bn1.weight',
            'Resnet.module.layer2.2.bn1.bias',
            'Resnet.module.layer2.2.conv2.weight', 'Resnet.module.layer2.2.bn2.weight',
            'Resnet.module.layer2.2.bn2.bias',
            'Resnet.module.layer2.2.conv3.weight', 'Resnet.module.layer2.2.bn3.weight',
            'Resnet.module.layer2.2.bn3.bias',
            'Resnet.module.layer2.3.conv1.weight', 'Resnet.module.layer2.3.bn1.weight',
            'Resnet.module.layer2.3.bn1.bias',
            'Resnet.module.layer2.3.conv2.weight', 'Resnet.module.layer2.3.bn2.weight',
            'Resnet.module.layer2.3.bn2.bias',
            'Resnet.module.layer2.3.conv3.weight', 'Resnet.module.layer2.3.bn3.weight',
            'Resnet.module.layer2.3.bn3.bias',
            'Resnet.module.layer3.0.conv1.weight', 'Resnet.module.layer3.0.bn1.weight',
            'Resnet.module.layer3.0.bn1.bias',
            'Resnet.module.layer3.0.conv2.weight', 'Resnet.module.layer3.0.bn2.weight',
            'Resnet.module.layer3.0.bn2.bias',
            'Resnet.module.layer3.0.conv3.weight', 'Resnet.module.layer3.0.bn3.weight',
            'Resnet.module.layer3.0.bn3.bias',
            'Resnet.module.layer3.0.downsample.0.weight', 'Resnet.module.layer3.0.downsample.1.weight',
            'Resnet.module.layer3.0.downsample.1.bias',
            'Resnet.module.layer3.1.conv1.weight', 'Resnet.module.layer3.1.bn1.weight',
            'Resnet.module.layer3.1.bn1.bias',
            'Resnet.module.layer3.1.conv2.weight', 'Resnet.module.layer3.1.bn2.weight',
            'Resnet.module.layer3.1.bn2.bias',
            'Resnet.module.layer3.1.conv3.weight', 'Resnet.module.layer3.1.bn3.weight',
            'Resnet.module.layer3.1.bn3.bias',
            'Resnet.module.layer3.2.conv1.weight', 'Resnet.module.layer3.2.bn1.weight',
            'Resnet.module.layer3.2.bn1.bias',
            'Resnet.module.layer3.2.conv2.weight', 'Resnet.module.layer3.2.bn2.weight',
            'Resnet.module.layer3.2.bn2.bias',
            'Resnet.module.layer3.2.conv3.weight', 'Resnet.module.layer3.2.bn3.weight',
            'Resnet.module.layer3.2.bn3.bias',
            'Resnet.module.layer3.3.conv1.weight', 'Resnet.module.layer3.3.bn1.weight',
            'Resnet.module.layer3.3.bn1.bias',
            'Resnet.module.layer3.3.conv2.weight', 'Resnet.module.layer3.3.bn2.weight',
            'Resnet.module.layer3.3.bn2.bias',
            'Resnet.module.layer3.3.conv3.weight', 'Resnet.module.layer3.3.bn3.weight',
            'Resnet.module.layer3.3.bn3.bias',
            'Resnet.module.layer3.4.conv1.weight', 'Resnet.module.layer3.4.bn1.weight',
            'Resnet.module.layer3.4.bn1.bias',
            'Resnet.module.layer3.4.conv2.weight', 'Resnet.module.layer3.4.bn2.weight',
            'Resnet.module.layer3.4.bn2.bias',
            'Resnet.module.layer3.4.conv3.weight', 'Resnet.module.layer3.4.bn3.weight',
            'Resnet.module.layer3.4.bn3.bias',
            'Resnet.module.layer3.5.conv1.weight', 'Resnet.module.layer3.5.bn1.weight',
            'Resnet.module.layer3.5.bn1.bias',
            'Resnet.module.layer3.5.conv2.weight', 'Resnet.module.layer3.5.bn2.weight',
            'Resnet.module.layer3.5.bn2.bias',
            'Resnet.module.layer3.5.conv3.weight', 'Resnet.module.layer3.5.bn3.weight',
            'Resnet.module.layer3.5.bn3.bias',
            'Resnet.module.layer4.0.conv1.weight', 'Resnet.module.layer4.0.bn1.weight',
            'Resnet.module.layer4.0.bn1.bias',
            'Resnet.module.layer4.0.conv2.weight', 'Resnet.module.layer4.0.bn2.weight',
            'Resnet.module.layer4.0.bn2.bias',
            'Resnet.module.layer4.0.conv3.weight', 'Resnet.module.layer4.0.bn3.weight',
            'Resnet.module.layer4.0.bn3.bias',
            'Resnet.module.layer4.0.downsample.0.weight', 'Resnet.module.layer4.0.downsample.1.weight',
            'Resnet.module.layer4.0.downsample.1.bias',
            'Resnet.module.layer4.1.conv1.weight', 'Resnet.module.layer4.1.bn1.weight',
            'Resnet.module.layer4.1.bn1.bias',
            'Resnet.module.layer4.1.conv2.weight', 'Resnet.module.layer4.1.bn2.weight',
            'Resnet.module.layer4.1.bn2.bias',
            'Resnet.module.layer4.1.conv3.weight', 'Resnet.module.layer4.1.bn3.weight',
            'Resnet.module.layer4.1.bn3.bias',
            'Resnet.module.layer4.2.conv1.weight', 'Resnet.module.layer4.2.bn1.weight',
            'Resnet.module.layer4.2.bn1.bias',
            'Resnet.module.layer4.2.conv2.weight', 'Resnet.module.layer4.2.bn2.weight',
            'Resnet.module.layer4.2.bn2.bias',
            'Resnet.module.layer4.2.conv3.weight', 'Resnet.module.layer4.2.bn3.weight',
            'Resnet.module.layer4.2.bn3.bias',
            'Resnet.module.liner.weight', 'Resnet.module.liner.bias'
        ]

        for name, para in model_vit.named_parameters():
            # 除head, pre_logits外，其他权重全部冻结False表示冻结
            # if "head" not in name and "pre_logits" not in name:
            # if name not in miss_list:
            if name in miss_list:  # 1.27测试
                para.requires_grad_(False)
                print("freeze {}".format(name))

            else:
                para.requires_grad_(True)
                print("training {}".format(name))


    # ======================================= Model Dnn=======================================================
    # model_dnn = points_to_dnn(args).to(device)
    # ================================================================================================
    pg = [p for p in model_vit.parameters() if p.requires_grad]
    optimizer = optim.SGD(pg, lr=args.lr, momentum=0.9, weight_decay=5E-5)
    # Scheduler https://arxiv.org/pdf/1812.01187.pdf
    lf = lambda x: ((1 + math.cos(x * math.pi / args.epochs)) / 2) * (1 - args.lrf) + args.lrf  # cosine
    scheduler = lr_scheduler.LambdaLR(optimizer, lr_lambda=lf)

    # 创建train_acc.csv和var_acc.csv文件，记录loss和accuracy
    df = pd.DataFrame(columns=['time', 'step', 'train Loss', 'training accuracy', 'val Loss', 'val accuracy'])  # 列名
    df.to_csv("./val_acc.csv", index=False)  # 路径可以根据需要更改

    for epoch in range(args.epochs):
        sheet2.write(epoch + 1, 0, epoch + 1)
        sheet2.write(epoch + 1, 5, str(optimizer.state_dict()['param_groups'][0]['lr']))

        # train
        # model_1,model_2,不冻结权重，依次获取值，
        train_loss, train_acc = train_one_epoch(model=model_vit,
                                                optimizer=optimizer,
                                                data_loader=train_dataloader,
                                                device=device,
                                                epoch=epoch)

        scheduler.step()

        sheet2.write(epoch + 1, 1, str(train_loss))
        sheet2.write(epoch + 1, 2, str(train_acc))

        # validate
        val_loss, val_acc = evaluate(model=model_vit,
                                     data_loader=val_dataloader,
                                     device=device,
                                     epoch=epoch)

        sheet2.write(epoch + 1, 3, str(val_loss))
        sheet2.write(epoch + 1, 4, str(val_acc))

        # time = "%s" % datetime.now()  # 获取当前时间
        # step = "Step[%d]" % epoch
        # list = [time, step, train_loss, train_acc, val_loss, val_acc]
        # data = pd.DataFrame([list])
        # data.to_csv('./val_acc.csv', mode='a', header=False, index=False)  # mode设为a,就可以向csv文件追加数据了

        tags = ["train_loss", "train_acc", "val_loss", "val_acc", "learning_rate"]
        tb_writer.add_scalar(tags[0], train_loss, epoch)
        tb_writer.add_scalar(tags[1], train_acc, epoch)
        tb_writer.add_scalar(tags[2], val_loss, epoch)
        tb_writer.add_scalar(tags[3], val_acc, epoch)
        tb_writer.add_scalar(tags[4], optimizer.param_groups[0]["lr"], epoch)

        if val_acc > best_acc:
            best_acc = val_acc
            # torch.save(model_vit.state_dict(), "./weights/best_model2.pth")  vit 训练权重
            torch.save(model_vit.state_dict(), "./weights/best_model_FABO_Cro_enh_16_1_bata.pth")  # dnn训练权重
            # torch.save(model.state_dict(), "./weights/model-{}.pth".format(epoch))

        sheet2.write(1, 6, str(best_acc))
        book.save('./Train_FABO_Cro_enh_16_1_bata.xlsx')  # vit
        # book.save('./Train_dnn_data.xlsx')   # dnn
        print("The Best Acc = : {:.4f}".format(best_acc))

        # torch.save(model_vit.state_dict(), "./weights/train_lest_FABO_Cro_rnh_4+6.pth")

        output_file_word = f"model_layers_and_params_epoch_{epoch + 1}.docx"
        output_file_excel = f"model_layers_and_params_epoch_{epoch + 1}.xlsx"
        save_model_layers_and_params_to_word(model_vit, epoch + 1, output_file_word)
        save_model_layers_and_params_to_excel(model_vit, epoch + 1, output_file_excel)
        print(f"Model layers and parameters have been saved to {output_file_word} and {output_file_excel}")


if __name__ == '__main__':
    parser = argparse.ArgumentParser()
    parser.add_argument('--num_classes', type=int, default=11)
    parser.add_argument('--epochs', type=int, default=70)
    parser.add_argument('--batch_size', type=int, default=1)
    parser.add_argument('--lr', type=float, default=0.001)
    parser.add_argument('--lrf', type=float, default=0.01)
    # ============================================================
    parser.add_argument('--use_fusion', action="store_true", dest="use_fusion", default=False, help='use images and points fusion')
    parser.add_argument('--use_branch', action="store_true", dest="use_branch", default="images", help='use branch: images,face_points, body_points')
    parser.add_argument('--use_points_body', action="store_true", dest="use_points_body", default=False, help='use points body for dnn')
    parser.add_argument('--use_points_face', action="store_true", dest="use_points_face", default=False, help='use points face for dnn')
    # ============================================================
    parser.add_argument('--first_layer_size', type=int, default=768)
    parser.add_argument('--confidence_threshold', type=float, default=0.1)
    parser.add_argument('--top_N_frames', type=float, default=10)
    # ============================================================
 
    # 数据集所在根目录
    # CAER-S    D:/datasets/CAER-S
    # CAER      D:datasets/CAER
    # 这里default只需要指向解压后的数据集目录就可以，名字太长可以修改
    # parser.add_argument('--data_path', type=str, default="/home/ubuntu/nas/FB-GER/FABO/")
    parser.add_argument('--data_path', type=str, default="./FABO/")
    parser.add_argument('--model_name', default='', help='create model name')

    # todo 预训练权重路径，如果不想载入就设置为空字符
    # 这里default要指向预训练权重路径
    parser.add_argument('--weights', type=str, default='./weights/best_model_FABO_Cro_att.pth', help='initial weights path')
    # 是否冻结权重
    parser.add_argument('--freeze_layers', type=bool, default=False)
    parser.add_argument('--device', default='cuda:0', help='device id (i.e. 0 or 0,1 or cpu)')
    opt = parser.parse_args()
    main(opt)
    
    # data = xlrd.open_workbook('Train_data.xlsx')
    # # sheet_names = data.sheet_names()
    # # print(sheet_names)
    #
    # sh = data.sheet_by_name(u'Train_data')
    # for rownum in range(sh.nrows):
    #     print(sh.row_values(rownum))
